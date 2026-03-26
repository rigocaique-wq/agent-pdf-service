from datetime import datetime
from pathlib import Path
import contextlib
import html
import re
from io import BytesIO

from fastapi import FastAPI
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse
from mcp.server.fastmcp import FastMCP
from mcp.server.transport_security import TransportSecuritySettings

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PUBLIC_BASE_URL = "https://agent-pdf-service.onrender.com"

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = Path("/tmp/generated_docs")
HEADER_IMAGE = BASE_DIR / "header-doc.png"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

transport_security = TransportSecuritySettings(
    allowed_hosts=[
        "127.0.0.1:*",
        "localhost:*",
        "[::1]:*",
        "agent-pdf-service.onrender.com",
        "agent-pdf-service.onrender.com:*",
    ]
)

mcp = FastMCP(
    "render-word-server",
    json_response=True,
    transport_security=transport_security,
)

# Faz o endpoint MCP responder diretamente em /mcp/ quando montado em /mcp
mcp.settings.streamable_http_path = "/"


@mcp.tool()
def ping() -> str:
    """Simple connectivity test."""
    return "MCP server is connected and working."


def normalize_title(title: str) -> str:
    title = (title or "Project Documentation").strip()
    return html.escape(title)


def normalize_markdown(content: str) -> str:
    content = (content or "").strip()

    # Remove blocos ```markdown ... ``` se vierem do modelo
    content = re.sub(r"^```(?:markdown|md)?\s*", "", content, flags=re.IGNORECASE)
    content = re.sub(r"\s*```$", "", content)

    return content.strip()


def remove_duplicated_title_from_content(title: str, content: str) -> str:
    """
    Remove um H1 inicial duplicado do markdown quando ele repete o mesmo title.
    Exemplo:
    title = "My Doc"
    content começa com "# My Doc"
    """
    normalized_title = (title or "").strip().lower()
    lines = content.splitlines()

    if not lines:
        return content

    first_non_empty_idx = None
    for i, line in enumerate(lines):
        if line.strip():
            first_non_empty_idx = i
            break

    if first_non_empty_idx is None:
        return content

    first_line = lines[first_non_empty_idx].strip()

    if first_line.startswith("# "):
        heading_text = first_line[2:].strip().lower()
        if heading_text == normalized_title:
            del lines[first_non_empty_idx]

            while first_non_empty_idx < len(lines) and not lines[first_non_empty_idx].strip():
                del lines[first_non_empty_idx]

    return "\n".join(lines).strip()


def markdown_to_html(content: str) -> str:
    normalized = normalize_markdown(content)
    return markdown.markdown(
        normalized,
        extensions=["extra", "sane_lists", "nl2br", "tables", "fenced_code"],
        output_format="html5",
    )


def set_cell_text(cell, text: str):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")
    run.font.size = Pt(10.5)


def add_page_number(paragraph):
    """
    Adiciona campo dinâmico de número da página no rodapé do Word.
    """
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def ensure_custom_styles(doc: Document):
    styles = doc.styles

    if "CustomTitle" not in styles:
        style = styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style.font.size = Pt(20)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x35, 0x63, 0xB2)

    if "CustomHeading2" not in styles:
        style = styles.add_style("CustomHeading2", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Georgia"
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x17, 0x3A, 0x70)

    if "CustomHeading3" not in styles:
        style = styles.add_style("CustomHeading3", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Georgia"
        style.font.size = Pt(11)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x17, 0x3A, 0x70)

    if "CustomCode" not in styles:
        style = styles.add_style("CustomCode", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(9.5)


def add_inline_content(paragraph, node):
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
        return

    if not isinstance(node, Tag):
        return

    if node.name in ["strong", "b"]:
        run = paragraph.add_run(node.get_text())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x3D, 0x4F, 0x68)
    elif node.name in ["em", "i"]:
        run = paragraph.add_run(node.get_text())
        run.italic = True
        run.font.size = Pt(11)
    elif node.name == "code":
        run = paragraph.add_run(node.get_text())
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
    elif node.name == "a":
        run = paragraph.add_run(node.get_text())
        run.underline = True
        run.font.color.rgb = RGBColor(0x7B, 0x60, 0xC9)
        run.font.size = Pt(11)
    elif node.name == "br":
        paragraph.add_run("\n")
    else:
        for child in node.children:
            add_inline_content(paragraph, child)


def add_html_block_to_doc(doc: Document, node, list_level=0):
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(11)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "h1":
        p = doc.add_paragraph(style="CustomTitle")
        p.paragraph_format.space_after = Pt(18)
        p.add_run(node.get_text(strip=True))

    elif node.name == "h2":
        p = doc.add_paragraph(style="CustomHeading2")
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        p.add_run(node.get_text(strip=True))

    elif node.name == "h3":
        p = doc.add_paragraph(style="CustomHeading3")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(node.get_text(strip=True))

    elif node.name == "p":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        for child in node.children:
            add_inline_content(p, child)

    elif node.name == "ul":
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.15 * list_level)
            for child in li.children:
                if isinstance(child, Tag) and child.name in ["ul", "ol"]:
                    add_html_block_to_doc(doc, child, list_level + 1)
                else:
                    add_inline_content(p, child)

    elif node.name == "ol":
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.15 * list_level)
            for child in li.children:
                if isinstance(child, Tag) and child.name in ["ul", "ol"]:
                    add_html_block_to_doc(doc, child, list_level + 1)
                else:
                    add_inline_content(p, child)

    elif node.name == "blockquote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(node.get_text(" ", strip=True))
        run.italic = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x51, 0x64, 0x7F)

    elif node.name == "pre":
        p = doc.add_paragraph(style="CustomCode")
        p.paragraph_format.space_after = Pt(8)
        p.add_run(node.get_text())

    elif node.name == "hr":
        p = doc.add_paragraph()
        p.add_run("—" * 30)

    elif node.name == "table":
        rows = node.find_all("tr")
        if rows:
            max_cols = max(len(r.find_all(["th", "td"], recursive=False)) for r in rows)
            table = doc.add_table(rows=0, cols=max_cols)
            table.style = "Table Grid"

            for row_node in rows:
                row_cells = row_node.find_all(["th", "td"], recursive=False)
                row = table.add_row().cells
                for i, cell_node in enumerate(row_cells):
                    set_cell_text(row[i], cell_node.get_text(" ", strip=True))

    else:
        for child in node.children:
            add_html_block_to_doc(doc, child, list_level=list_level)


def build_word_document(title: str, html_body: str) -> Document:
    doc = Document()
    ensure_custom_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Header com imagem, se existir
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if HEADER_IMAGE.exists():
        try:
            run = header_p.add_run()
            run.add_picture(str(HEADER_IMAGE), width=Inches(6.8))
        except Exception:
            # Se der problema com a imagem, segue sem quebrar
            pass

    # Footer
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_run = footer_p.add_run(
        "3301 N University Drive, Suite 400, Coral Springs, FL 33065.   |   "
        "www.datameaning.com   |   "
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xA4, 0xB2, 0xC7)

    add_page_number(footer_p)

    # Título principal
    title_p = doc.add_paragraph(style="CustomTitle")
    title_p.paragraph_format.space_after = Pt(18)
    title_p.add_run(title)

    soup = BeautifulSoup(html_body, "html.parser")

    for node in soup.contents:
        add_html_block_to_doc(doc, node)

    return doc


@mcp.tool()
def generate_word(title: str, content: str) -> dict:
    """Generate a styled Word document from markdown content and return a download URL."""
    safe_timestamp = str(datetime.now().timestamp()).replace(".", "_")
    file_name = f"document_{safe_timestamp}.docx"
    file_path = OUTPUT_DIR / file_name

    safe_title = html.unescape(normalize_title(title))

    cleaned_content = normalize_markdown(content)
    cleaned_content = remove_duplicated_title_from_content(title, cleaned_content)

    body_html = markdown_to_html(cleaned_content)

    doc = build_word_document(safe_title, body_html)
    doc.save(str(file_path))

    download_url = f"{PUBLIC_BASE_URL}/files/{file_name}"

    return {
        "content": [
            {
                "type": "text",
                "text": f"Word document generated successfully.\nDownload it here: {download_url}"
            }
        ],
        "structuredContent": {
            "status": "success",
            "filename": file_name,
            "download_url": download_url
        }
    }


@contextlib.asynccontextmanager
async def lifespan(app: FastAPI):
    async with mcp.session_manager.run():
        yield


app = FastAPI(lifespan=lifespan)


@app.get("/")
def health():
    return {"message": "Server is running"}


# Redireciona /mcp -> /mcp/ preservando o método HTTP
@app.api_route("/mcp", methods=["GET", "POST", "HEAD", "OPTIONS"])
async def mcp_redirect():
    return RedirectResponse(url="/mcp/", status_code=307)


@app.get("/files/{filename}")
async def serve_file(filename: str):
    file_path = OUTPUT_DIR / filename

    if not file_path.exists():
        return JSONResponse({"error": "File not found"}, status_code=404)

    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


app.mount("/mcp", mcp.streamable_http_app())
